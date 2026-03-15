#!/usr/bin/env python3
"""
Full manuscript — Lancet style
Maternal Vitamin D and Fetal Growth / Adiposity: SR & Meta-Analysis
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os, re

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin = Inches(1.0)
    s.bottom_margin = Inches(1.0)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 2.0

NAVY = RGBColor(0x1A, 0x2E, 0x4A)
GREY = RGBColor(0x66, 0x66, 0x66)


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 2 else RGBColor(0x33, 0x33, 0x33)
        run.font.name = 'Times New Roman'
    return h


def para(text, bold=False, italic=False, size=12, align=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    return p


def para_mixed(parts, indent=False):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = bold
        run.italic = italic
    return p


# =============================================================================
# TITLE PAGE
# =============================================================================
para("RESEARCH ARTICLE", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run(
    "Association between maternal vitamin D levels in pregnancy "
    "and fetal growth and adiposity: a systematic review and meta-analysis"
)
r.bold = True
r.font.size = Pt(16)
r.font.name = 'Times New Roman'

doc.add_paragraph()
para("Isabel O. Almeida, Audencio Victor", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("Corresponding author: Audencio Victor (audenciovictor@gmail.com)",
     italic=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("PROSPERO registration: CRD42024590338", italic=True, size=11,
     align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# =============================================================================
# ABSTRACT
# =============================================================================
heading("Abstract", 1)

sections_abstract = [
    ("Background", (
        "Vitamin D [25(OH)D] is essential for fetal skeletal development, calcium homeostasis, "
        "and metabolic programming during pregnancy. Although maternal vitamin D deficiency is "
        "highly prevalent worldwide, the association between 25(OH)D concentrations and fetal "
        "growth outcomes remains inconsistent. We aimed to synthesise the available evidence "
        "on the relationship between maternal vitamin D status and fetal growth and adiposity."
    )),
    ("Methods", (
        "We conducted a systematic review and meta-analysis following PRISMA 2020 guidelines. "
        "Six databases (PubMed, EMBASE, Scopus, LILACS, SciELO, Web of Science) were searched "
        "up to October 2024. Observational studies and randomised controlled trials (RCTs) "
        "evaluating maternal serum 25(OH)D and fetal growth outcomes were included. Quality was "
        "assessed using JBI Critical Appraisal, Newcastle\u2013Ottawa Scale (NOS), and Cochrane "
        "RoB 2. Meta-analyses used fixed-effects (I\u00b2 < 50%) or random-effects (REML; "
        "I\u00b2 \u2265 50%) models. Publication bias was assessed via Egger\u2019s test and "
        "Duval & Tweedie trim-and-fill. Certainty of evidence was evaluated using GRADE."
    )),
    ("Findings", (
        "26 studies (24 observational, 2 RCTs; 26,542 participants; 16 countries) were included. "
        "Maternal vitamin D deficiency was significantly associated with increased risk of "
        "small-for-gestational-age or fetal growth restriction (SGA/FGR; pooled OR 1\u00b795 "
        "[95% CI 1\u00b747\u20132\u00b759]; I\u00b2 = 0%; k = 4; N = 9033; fixed-effects). "
        "For fetal biometry and estimated fetal weight, the pooled OR was 1\u00b716 "
        "(1\u00b709\u20131\u00b723; I\u00b2 = 33%; k = 4; N = 26,542; GRADE: high certainty). "
        "In RCTs, vitamin D supplementation significantly increased femur length "
        "(MD +1\u00b798 mm [1\u00b728\u20132\u00b768]) and humerus length "
        "(MD +1\u00b739 mm [0\u00b772\u20132\u00b706]). No significant association "
        "was found with fetal adiposity."
    )),
    ("Interpretation", (
        "Maternal vitamin D deficiency is consistently associated with impaired fetal growth, "
        "particularly SGA/FGR. Routine screening and correction of deficiency during pregnancy "
        "may improve fetal growth outcomes. Large-scale RCTs are warranted to establish "
        "causality and optimal supplementation strategies."
    )),
]

for label, text in sections_abstract:
    para_mixed([(label + " ", True, False), (text, False, False)])

para_mixed([("Keywords: ", True, False),
            ("vitamin D, 25-hydroxyvitamin D, pregnancy, fetal growth, "
             "small-for-gestational-age, fetal growth restriction, adiposity, "
             "meta-analysis, systematic review", False, False)])

doc.add_page_break()

# =============================================================================
# INTRODUCTION
# =============================================================================
heading("Introduction", 1)

para(
    "Vitamin D is a fat-soluble secosteroid that plays a central role in calcium and "
    "phosphorus homeostasis, immune regulation, cell proliferation, and gene "
    "transcription (1,2,4). The primary circulating form, 25-hydroxyvitamin D "
    "[25(OH)D], is synthesised in the liver from cholecalciferol (vitamin D3) obtained "
    "through cutaneous photosynthesis or dietary intake, and is subsequently converted to "
    "its biologically active form, 1,25-dihydroxyvitamin D [1,25(OH)2D3], in the "
    "kidneys and placenta (1).",
    indent=True
)

para(
    "During pregnancy, maternal requirements for vitamin D increase substantially to support "
    "fetal skeletal development, calcium transfer across the placenta, and metabolic "
    "programming (2,3). Maternal 25(OH)D crosses the placenta and serves as the "
    "primary source of vitamin D for the developing fetus, as fetal renal 1\u03b1-hydroxylase "
    "activity is limited (3). Consequently, maternal deficiency\u2014defined as serum "
    "25(OH)D below 50 nmol/L (20 ng/mL)\u2014directly affects fetal vitamin D status and may "
    "impair intrauterine growth (4).",
    indent=True
)

para(
    "The global prevalence of vitamin D deficiency among pregnant women is alarmingly high, "
    "estimated at 47\u00b79% in a pooled analysis of 7\u00b79 million participants (5). "
    "This burden is particularly pronounced in regions with limited sun exposure, high "
    "skin melanin content, and inadequate dietary supplementation (5,6). Given the "
    "established role of vitamin D in bone metabolism and placental function, several studies "
    "have investigated whether maternal deficiency is associated with adverse fetal growth "
    "outcomes, including small-for-gestational-age (SGA) birth, fetal growth restriction "
    "(FGR), altered fetal biometry parameters, and fetal adiposity (7\u201310).",
    indent=True
)

para(
    "However, the existing evidence remains inconsistent. While some studies have reported "
    "significant associations between low maternal 25(OH)D and increased risk of SGA or "
    "impaired fetal biometry (8,11,12), others have found no "
    "significant relationship (13,14). This heterogeneity may be partly "
    "explained by differences in study design, sample size, laboratory methods for 25(OH)D "
    "quantification, timing of blood sampling during pregnancy, and the definition of "
    "vitamin D deficiency thresholds (15).",
    indent=True
)

para(
    "Only two previous systematic reviews have specifically addressed the relationship "
    "between maternal vitamin D and fetal biometry or SGA risk in observational studies and "
    "RCTs, respectively. However, both were conducted before 2019 and did not include "
    "recent large-scale cohort studies that have since contributed substantially to the "
    "evidence base. Furthermore, no prior review has comprehensively evaluated fetal "
    "adiposity as an outcome in relation to maternal vitamin D status.",
    indent=True
)

para(
    "In view of these gaps, we conducted a systematic review and meta-analysis to synthesise "
    "the available evidence on the association between maternal 25(OH)D concentrations "
    "during pregnancy and fetal growth and adiposity outcomes.",
    indent=True
)

doc.add_page_break()

# =============================================================================
# METHODS
# =============================================================================
heading("Methods", 1)

heading("Study design and registration", 2)
para(
    "This systematic review was conducted in accordance with the Preferred Reporting Items "
    "for Systematic Reviews and Meta-Analyses (PRISMA) 2020 guidelines (16) and was "
    "prospectively registered at the International Prospective Register of Systematic Reviews "
    "(PROSPERO; CRD42024590338)."
)

heading("Eligibility criteria", 2)
para(
    "We included observational studies (cohort, case-control, cross-sectional) and randomised "
    "controlled trials (RCTs) that: (1) enrolled pregnant women; (2) measured maternal serum "
    "25(OH)D concentrations; and (3) reported at least one fetal growth or adiposity outcome, "
    "including SGA, FGR, intrauterine growth restriction (IUGR), fetal biometry parameters "
    "(femur length [FL], abdominal circumference [AC], biparietal diameter [BPD], head "
    "circumference [HC]), estimated fetal weight (EFW), or fetal adiposity measures. "
    "Studies involving multiple pregnancies, reviews, editorials, conference abstracts, "
    "case reports, and animal studies were excluded."
)

heading("Information sources and search strategy", 2)
para(
    "Systematic searches were conducted in six electronic databases: PubMed, EMBASE, "
    "Scopus, LILACS, SciELO, and Web of Science. The search was performed from database "
    "inception to October 12, 2024. Search strategies combined Medical Subject Headings "
    "(MeSH) and free-text terms related to vitamin D (e.g., \u201c25-hydroxyvitamin D\u201d, "
    "\u201ccholecalciferol\u201d, \u201cergocalciferol\u201d) and fetal growth outcomes "
    "(e.g., \u201cfetal growth\u201d, \u201csmall for gestational age\u201d, "
    "\u201cbiometry\u201d, \u201cadiposity\u201d). No language restrictions were applied. "
    "Reference lists of included studies and relevant reviews were hand-searched. "
    "The search strategy was developed using the PRESS guidelines. (17)"
)

heading("Study selection", 2)
para(
    "Two reviewers (AV and IOA) independently screened titles and abstracts in phase 1, "
    "followed by full-text assessment in phase 2. Disagreements were resolved by consensus. "
    "Zotero (version 7.0.21; Digital Scholar) was used for reference management."
)

heading("Data extraction", 2)
para(
    "Data were extracted by two independent reviewers using a standardised form. The following "
    "variables were recorded: first author, year, country, study design, sample size, "
    "participant characteristics (age, gestational age, BMI), method of 25(OH)D measurement, "
    "mean or median 25(OH)D concentration, vitamin D deficiency threshold, trimester of "
    "sampling, fetal outcomes assessed, effect estimates (odds ratios, risk ratios, mean "
    "differences, correlation coefficients), and statistical methods used."
)

heading("Quality assessment", 2)
para(
    "Methodological quality of observational studies was assessed using the Joanna Briggs "
    "Institute (JBI) Critical Appraisal tools (18) and the Newcastle\u2013Ottawa Scale "
    "(NOS). (19) The JBI tools consist of domain-specific checklists (8\u201311 items) "
    "evaluated as \u201cYes\u201d, \u201cNo\u201d, or \u201cUnclear\u201d. The NOS was scored "
    "on a 0\u20139 scale, with scores \u22657 indicating high quality. Risk of bias in RCTs "
    "was assessed using the Cochrane Risk of Bias 2 (RoB 2) tool. (20)"
)

heading("Statistical analysis", 2)
para(
    "Meta-analyses were performed using R (version 4.3) with the meta and metafor packages. "
    "For dichotomous outcomes (SGA/FGR, biometry abnormalities), odds ratios (OR) with 95% "
    "confidence intervals (CI) were pooled using inverse-variance weighting. Where original "
    "studies reported protective ORs (i.e., higher vitamin D associated with lower risk), "
    "effect estimates were inverted to harmonise the direction of association such that "
    "vitamin D deficiency corresponded to increased risk."
)
para(
    "For continuous outcomes from RCTs, mean differences (MD) with 95% CIs were calculated. "
    "An adaptive model selection approach was used: fixed-effects (Mantel\u2013Haenszel) "
    "models were applied when heterogeneity was low (I\u00b2 < 50%), and random-effects "
    "models with restricted maximum likelihood (REML) estimation of between-study variance "
    "(\u03c4\u00b2) were used when I\u00b2 \u2265 50%.",
    indent=True
)
para(
    "Heterogeneity was quantified using Cochran\u2019s Q statistic and the I\u00b2 index. "
    "Sensitivity analyses included leave-one-out analysis and restriction to prospective "
    "cohort studies only. Subgroup analysis was performed by vitamin D deficiency threshold "
    "(< 25, \u2264 30, \u2264 50 nmol/L). Publication bias was assessed using Egger\u2019s "
    "linear regression test, Begg\u2019s rank correlation test, and Duval & Tweedie "
    "trim-and-fill analysis. The clinical relevance of the pooled OR was expressed as "
    "risk difference (RD) and number needed to harm (NNH), assuming a baseline SGA "
    "prevalence of 8%. Certainty of evidence was evaluated for each outcome using the "
    "Grading of Recommendations Assessment, Development and Evaluation (GRADE) "
    "framework. (21)",
    indent=True
)

heading("Role of the funding source", 2)
para("No external funding was received for this study.")

doc.add_page_break()

# =============================================================================
# RESULTS
# =============================================================================
heading("Results", 1)

heading("Study selection", 2)
para(
    "The systematic search identified 3135 records across six databases. After removal of "
    "duplicates, 3412 unique records were screened by title and abstract (phase 1), of which "
    "30 were assessed in full text (phase 2). Four studies were excluded: two did not measure "
    "fetal growth outcomes, one included only postpartum vitamin D assessment, and one was a "
    "duplicate publication. A total of 26 studies met the eligibility criteria and were "
    "included in the qualitative synthesis. Of these, eight studies (four for SGA/FGR and "
    "four for biometry/EFW) provided sufficient data for quantitative meta-analysis (figure 1)."
)

heading("Study characteristics", 2)
para(
    "The 26 included studies comprised 24 observational studies (18 prospective or "
    "retrospective cohorts, 2 case-control, 2 cross-sectional, 2 population-based) and "
    "2 RCTs, enrolling a total of 26,542 participants (table 1). Studies were conducted in "
    "16 countries across four continents: Asia (12 studies, 46%), Europe (8, 31%), the "
    "Americas (4, 15%), and Africa (1, 4%). Sample sizes ranged from 48 to 10,913 "
    "participants. Publication dates spanned 2010 to 2025, with 14 studies (54%) published "
    "after 2020."
)
para(
    "Maternal 25(OH)D concentrations were reported in 22 studies, with a mean of "
    "53\u00b76 \u00b1 14\u00b77 nmol/L (median 50\u00b71 nmol/L; IQR 43\u00b76\u201365\u00b70). "
    "Concentrations varied by continent: studies from Asia reported the lowest mean values "
    "(48\u00b70 \u00b1 14\u00b71 nmol/L), while American studies reported the highest "
    "(61\u00b78 \u00b1 5\u00b77 nmol/L). Vitamin D was measured by immunoassay in 18 studies "
    "(69%; CLIA, RIA, ECLIA, ELISA, ECL) and by mass spectrometry in 8 studies (31%; "
    "HPLC-MS/MS or HPLC). The most commonly used deficiency threshold was "
    "\u226450 nmol/L (17 studies, 65%).",
    indent=True
)
para(
    "The outcomes assessed were diverse: 10 studies evaluated SGA, FGR, or IUGR; 8 assessed "
    "fetal biometry or EFW; 4 evaluated fetal bone parameters (FL, HL); 1 assessed fetal "
    "adiposity; and 3 evaluated combined or other growth-related outcomes. Of 26 studies, "
    "15 (58%) reported statistically significant associations between maternal vitamin D "
    "status and at least one fetal growth outcome.",
    indent=True
)

heading("Methodological quality", 2)
para(
    "Among the 24 observational studies, JBI Critical Appraisal scores ranged from 4 to 11 "
    "(out of 11), with most studies meeting 7 or more criteria. NOS scores ranged from 5 to "
    "9 (out of 9): 11 studies (46%) were rated as high quality (NOS \u22657), 11 (46%) as "
    "moderate (NOS 5\u20136), and 2 (8%) as lower quality (NOS < 5). Common methodological "
    "limitations included inadequate adjustment for confounders, incomplete reporting of "
    "follow-up, and lack of standardisation in vitamin D measurement methods (figure S1; "
    "table S3)."
)
para(
    "For the two RCTs, assessment using the Cochrane RoB 2 tool indicated that Vafaei et al. "
    "(2019) had a low overall risk of bias, while Srilekha et al. (2021) had a high risk of "
    "bias, primarily due to concerns about the randomisation process, lack of outcome "
    "assessor blinding, and unclear allocation concealment.",
    indent=True
)

heading("Primary meta-analysis: SGA/FGR", 2)
para(
    "Four studies (Miliku et al., 2016; Gernand et al., 2014 [two strata]; Beck et al., "
    "2025) involving 9033 participants provided sufficient data for pooled analysis of "
    "the association between maternal vitamin D deficiency and SGA/FGR risk (figure 2). "
    "All four studies used HPLC-MS/MS for 25(OH)D quantification (the reference standard "
    "method). The pooled odds ratio was 1\u00b795 (95% CI 1\u00b747\u20132\u00b759; "
    "I\u00b2 = 0%; Cochran\u2019s Q = 0\u00b779, df = 3, p = 0\u00b7853), using a "
    "fixed-effects model. All individual study estimates were directionally consistent, "
    "with ORs ranging from 1\u00b728 to 2\u00b717."
)
para(
    "The clinical significance was quantified assuming a baseline SGA prevalence of 8%: the "
    "risk difference was +7\u00b71%, corresponding to a number needed to harm (NNH) of 14. "
    "This indicates that for every 14 pregnancies with vitamin D deficiency, one additional "
    "case of SGA/FGR would be expected compared with vitamin D-sufficient pregnancies.",
    indent=True
)

heading("Sensitivity analyses for SGA/FGR", 2)
para(
    "Leave-one-out analysis confirmed the robustness of the pooled estimate: all leave-one-out "
    "ORs remained statistically significant and ranged from 1\u00b766 to 2\u00b710 "
    "(figure S2). When the analysis was restricted to prospective and observational cohort "
    "studies only (excluding case-control designs), the pooled OR was similar "
    "(OR 1\u00b793 [95% CI 1\u00b745\u20132\u00b757])."
)
para(
    "In a sensitivity analysis including Alimohammadi et al. (2020)\u2014a case-control "
    "study with an approximate confidence interval (OR 6\u00b781 [2\u00b720\u201321\u00b714])"
    "\u2014the pooled estimate shifted to OR 2\u00b746 (1\u00b745\u20134\u00b717) under a "
    "random-effects model (I\u00b2 = 52%), suggesting that this study was a potential outlier "
    "and appropriately excluded from the primary analysis.",
    indent=True
)
para(
    "Subgroup analysis by vitamin D deficiency threshold showed directionally consistent "
    "results: studies using a cutoff of < 25 nmol/L (k = 1; OR 2\u00b707), \u2264 30 nmol/L "
    "(k = 2; pooled OR 1\u00b795), and \u2264 50 nmol/L (k = 1; OR 1\u00b728) all favoured "
    "an increased risk with deficiency (figure S3).",
    indent=True
)

heading("Secondary meta-analysis: fetal biometry and EFW", 2)
para(
    "Four study estimates from two independent cohorts (Liu et al., 2020 [two strata]; "
    "Morales et al., 2015 [AC and EFW]) involving 26,542 participants were pooled for "
    "the biometry/EFW outcome (figure 3). The pooled OR was 1\u00b716 (95% CI "
    "1\u00b709\u20131\u00b723; I\u00b2 = 33%; Cochran\u2019s Q = 4\u00b745, df = 3, "
    "p = 0\u00b7217), using a fixed-effects model. The GRADE certainty of evidence was "
    "rated as high."
)
para(
    "Liu et al. (2020) reported that the combined effect of vitamin D deficiency and "
    "gestational diabetes mellitus (GDM) on excessive EFW was stronger (OR 1\u00b736 "
    "[1\u00b715\u20131\u00b762]) than vitamin D deficiency alone (OR 1\u00b711 "
    "[1\u00b702\u20131\u00b721]), suggesting a potential synergistic effect. "
    "Morales et al. (2015) found that each 10 ng/mL decrease in maternal 25(OH)D at "
    "10 weeks\u2019 gestation was associated with increased odds of AC < 10th centile "
    "(OR 1\u00b714 [1\u00b700\u20131\u00b731]) and EFW < 10th centile "
    "(OR 1\u00b718 [1\u00b703\u20131\u00b736]).",
    indent=True
)

heading("RCT continuous outcomes", 2)
para(
    "Two RCTs reported continuous fetal biometry outcomes with vitamin D supplementation "
    "(figure 4). Vafaei et al. (2019), a low risk-of-bias trial of 140 pregnant women "
    "receiving 1000 IU/day vitamin D, found significant increases in femur length "
    "(MD +1\u00b798 mm [95% CI 1\u00b728\u20132\u00b768]; p < 0\u00b7001) and humerus "
    "length (MD +1\u00b739 mm [0\u00b772\u20132\u00b706]; p < 0\u00b7001) at the second "
    "trimester."
)
para(
    "Srilekha et al. (2021), classified as high risk of bias, reported larger effects with "
    "supplementation: EFW +277 g (p < 0\u00b701), BPD +4\u00b777 mm, and FL +3\u00b703 mm. "
    "However, these findings should be interpreted with caution given the methodological "
    "limitations, including unclear allocation concealment and absence of outcome assessor "
    "blinding.",
    indent=True
)

heading("Publication bias", 2)
para(
    "Egger\u2019s linear regression test could not be reliably performed for the SGA/FGR "
    "meta-analysis due to the small number of studies (k = 4), which limits the statistical "
    "power of funnel plot asymmetry tests. Begg\u2019s rank correlation test was non-significant. "
    "Duval & Tweedie trim-and-fill analysis imputed no additional studies for SGA/FGR, "
    "indicating no evidence of funnel plot asymmetry (figure S4). For biometry/EFW, "
    "trim-and-fill analysis also imputed no studies, and the adjusted estimate remained "
    "unchanged."
)

heading("Narrative synthesis of remaining studies", 2)
para(
    "Eighteen studies could not be included in the meta-analysis due to heterogeneous "
    "outcome definitions, non-comparable effect measures, or insufficient reporting of "
    "quantitative data. Among these, notable findings included: Lee D.H. et al. (2015) "
    "reported a significant correlation between changes in maternal 25(OH)D and changes "
    "in BPD during pregnancy (r = 0\u00b714, p = 0\u00b703) using generalised estimating "
    "equations; Young et al. (2012) found a significant vitamin D\u2013calcium interaction "
    "effect on fetal femur z-score (\u03b2 = 0\u00b715, p < 0\u00b705); and Lee S.B. et al. "
    "(2023) reported that severe deficiency (< 25 nmol/L) was associated with developmental "
    "delay (aOR 4\u00b728 [1\u00b740\u201313\u00b705])."
)
para(
    "Regarding fetal adiposity, only Akita et al. (2025) and Morales et al. (2015) provided "
    "relevant data. Neither study found a significant association between maternal 25(OH)D "
    "and fetal adiposity measures, though both had limited sample sizes for this specific "
    "outcome. The GRADE certainty for adiposity was rated as low.",
    indent=True
)

heading("GRADE evidence profile", 2)
para(
    "The certainty of evidence varied by outcome (figure S5): high for biometry/EFW "
    "(consistent, precise, large sample), moderate for SGA/FGR (downgraded for indirectness "
    "due to varied deficiency thresholds), moderate for femur and humerus length from RCTs "
    "(single study), low for EFW/BPD from the high risk-of-bias RCT, and low for fetal "
    "adiposity (imprecision and limited data)."
)

doc.add_page_break()

# =============================================================================
# DISCUSSION
# =============================================================================
heading("Discussion", 1)

para(
    "This systematic review and meta-analysis, comprising 26 studies and 26,542 participants "
    "from 16 countries, provides robust evidence that maternal vitamin D deficiency during "
    "pregnancy is associated with impaired fetal growth. The primary meta-analysis "
    "demonstrated a nearly two-fold increased risk of SGA/FGR (OR 1\u00b795 "
    "[1\u00b747\u20132\u00b759]) among vitamin D-deficient pregnancies, with no "
    "heterogeneity (I\u00b2 = 0%) and consistent results across all sensitivity analyses. "
    "The secondary analysis confirmed a modest but significant association with altered "
    "fetal biometry and EFW (OR 1\u00b716 [1\u00b709\u20131\u00b723]; GRADE: high "
    "certainty).",
    indent=True
)

para(
    "Our findings for SGA/FGR are consistent with, and extend, previous meta-analyses. "
    "A 2018 meta-analysis of observational studies reported a pooled OR of 1\u00b755 "
    "(1\u00b730\u20131\u00b784) for the association between vitamin D deficiency and SGA, "
    "based on a larger but methodologically heterogeneous pool of studies. (22) "
    "Our estimate is higher (OR 1\u00b795), which may reflect the restriction to studies "
    "using HPLC-MS/MS\u2014the reference standard for 25(OH)D measurement\u2014in the "
    "primary analysis, thereby reducing measurement error and potential attenuation bias. "
    "A meta-analysis of RCTs similarly reported that vitamin D supplementation reduced "
    "SGA risk (RR 0\u00b772 [0\u00b757\u20130\u00b792]), (23) providing "
    "complementary evidence from an interventional perspective.",
    indent=True
)

para(
    "The biological plausibility of the observed association is well-supported. Vitamin D "
    "modulates placental calcium transport through regulation of calbindin-D and "
    "TRPV6 channels, both essential for fetal skeletal mineralisation. (24) "
    "Additionally, 25(OH)D exerts immunomodulatory effects at the maternal\u2013fetal "
    "interface, influencing trophoblast invasion, spiral artery remodelling, and "
    "anti-inflammatory cytokine production. (25) Deficiency may therefore impair "
    "both nutrient transfer and placental vascular development, leading to fetal growth "
    "restriction.",
    indent=True
)

para(
    "The finding that all studies included in the SGA/FGR meta-analysis used HPLC-MS/MS "
    "is noteworthy. Immunoassays, which were used in the majority of included studies "
    "(69%), are known to have variable cross-reactivity with vitamin D metabolites and "
    "lower accuracy compared with mass spectrometry-based methods. (26) This "
    "methodological difference may contribute to the inconsistent findings across "
    "observational studies and underscores the importance of standardised measurement "
    "approaches in future research.",
    indent=True
)

para(
    "The geographic diversity of included studies is both a strength and a source of "
    "heterogeneity. Asian studies, which comprised 46% of the evidence base, reported "
    "the lowest mean 25(OH)D concentrations (48\u00b70 nmol/L) and the highest proportion "
    "of significant associations. This geographic variation likely reflects differences in "
    "sun exposure, skin pigmentation, dietary habits, and supplementation practices.\u2075 "
    "The subgroup analysis by deficiency threshold demonstrated consistent results "
    "regardless of the cutoff used, suggesting that the association is robust across "
    "different definitions of deficiency.",
    indent=True
)

para(
    "The RCT evidence, while limited to two trials, provides preliminary support for a "
    "causal relationship. The significant increase in femur and humerus length with "
    "supplementation reported by Vafaei et al. (2019)\u2014a well-conducted trial with "
    "low risk of bias\u2014is biologically coherent with the role of vitamin D in "
    "endochondral ossification. (27) However, the larger effect sizes reported "
    "by Srilekha et al. (2021) should be interpreted cautiously given the high risk of "
    "bias. The paucity of high-quality RCTs remains a significant limitation of the "
    "evidence base.",
    indent=True
)

para(
    "No significant association was found between maternal 25(OH)D and fetal adiposity, "
    "though this conclusion is limited by the small number of studies addressing this "
    "outcome (k = 2) and the low GRADE certainty. Animal studies have suggested that "
    "vitamin D may influence adipogenesis through effects on PPAR\u03b3 and Wnt signalling "
    "pathways, (28) but human evidence remains insufficient to confirm this "
    "hypothesis in the fetal compartment.",
    indent=True
)

heading("Strengths and limitations", 2)
para(
    "This review has several strengths. First, the comprehensive search across six "
    "databases identified a large and geographically diverse evidence base. Second, "
    "the adaptive model selection approach ensured appropriate handling of heterogeneity. "
    "Third, the use of multiple sensitivity analyses, subgroup analyses, and trim-and-fill "
    "analysis strengthened confidence in the pooled estimates. Fourth, the dual quality "
    "assessment (JBI + NOS for observational studies; RoB 2 for RCTs) and GRADE evaluation "
    "provided transparent assessment of evidence certainty.",
    indent=True
)
para(
    "Limitations include the small number of studies eligible for meta-analysis (k = 4 per "
    "outcome), which limited the power of publication bias tests and subgroup analyses. "
    "The heterogeneity in vitamin D deficiency thresholds and timing of 25(OH)D measurement "
    "across studies may have introduced variability. Only two RCTs were available, one with "
    "high risk of bias, precluding definitive causal conclusions. Finally, most observational "
    "studies did not adequately adjust for key confounders such as maternal BMI, ethnicity, "
    "season of blood sampling, and dietary calcium intake.",
    indent=True
)

doc.add_page_break()

# =============================================================================
# CONCLUSION
# =============================================================================
heading("Conclusion", 1)
para(
    "Maternal vitamin D deficiency during pregnancy is significantly and consistently "
    "associated with impaired fetal growth, particularly increased risk of SGA/FGR "
    "(OR 1\u00b795) and altered fetal biometry. Preliminary RCT evidence suggests that "
    "supplementation may improve fetal bone growth, though the evidence remains limited. "
    "These findings support the rationale for routine screening and correction of vitamin D "
    "deficiency in pregnancy as a potential strategy to optimise fetal growth outcomes. "
    "Large-scale, well-designed RCTs with standardised supplementation protocols, "
    "consistent deficiency thresholds, and fetal adiposity as a primary endpoint are "
    "urgently needed to establish causality and inform clinical guidelines."
)

doc.add_page_break()

# =============================================================================
# REFERENCES
# =============================================================================
heading("References", 1)

refs = [
    "1. Bikle DD, Christakos S. New aspects of vitamin D metabolism and action\u2014addressing the skin as source and target. Nat Rev Endocrinol 2020; 16: 234\u201352.",
    "2. Hossein-nezhad A, Holick MF. Vitamin D for health: a global perspective. Mayo Clin Proc 2013; 88: 720\u201355.",
    "3. Palaniswamy S, Williams D, J\u00e4rvelin MR, Sebert S. Vitamin D and the promotion of long-term metabolic health from a programming perspective. Nutr Metab Insights 2015; 8s1: NMI.S29526.",
    "4. Agarwal S, Kovilam O, Agrawal DK. Vitamin D and its impact on maternal-fetal outcomes in pregnancy: a critical review. Crit Rev Food Sci Nutr 2018; 58: 755\u201369.",
    "5. Cui A, Zhang T, Xiao P, et al. Global and regional prevalence of vitamin D deficiency in population-based studies from 2000 to 2022: a pooled analysis of 7.9 million participants. Front Nutr 2023; 10: 1070808.",
    "6. Dunlop E, Pham NM, Van Hoang D, et al. A systematic review and meta-analysis of circulating 25-hydroxyvitamin D concentration and vitamin D status in pregnancy. J Nutr 2025 (in press).",
    "7. Miliku K, Vinkhuyzen A, Blanken LM, et al. Maternal vitamin D concentrations during pregnancy, fetal growth patterns, and risks of adverse birth outcomes. Am J Clin Nutr 2016; 103: 1514\u201322.",
    "8. Gernand AD, Simhan HN, Caritis S, Bodnar LM. Maternal vitamin D status and small-for-gestational-age offspring in women at high risk for preeclampsia. Obstet Gynecol 2014; 123: 40\u20138.",
    "9. Liu Z, Liu H, Xu X, et al. Combined effect of maternal vitamin D deficiency and gestational diabetes mellitus on trajectories of ultrasound-measured fetal growth. Am J Clin Nutr 2020; 111: 378\u201388.",
    "10. Morales E, Rodriguez A, Valvi D, et al. Deficit of vitamin D in pregnancy and growth and overweight in the offspring. Int J Obes 2015; 39: 61\u20138.",
    "11. Beck C, Blue NR, Silver RM, et al. Maternal vitamin D status, fetal growth patterns, and adverse pregnancy outcomes in a multisite prospective pregnancy cohort. Am J Clin Nutr 2025; 121: 532\u201342.",
    "12. Tosun G, Akar AN, Burkankulu D, Ceyhan V, Aydin E. Effect of combination of uterine artery doppler and vitamin D level on perinatal outcomes in second trimester pregnant women. J Perinat Med 2025; 53: 48\u201357.",
    "13. Wierzejska R, Jarosz M, Sawicki W, Bachanek M, Siuba-Strzelinska M. Vitamin D concentration in maternal and umbilical cord blood by season. Int J Environ Res Public Health 2020; 17: 5689.",
    "14. Palmrich P, Thajer A, Schirwani N, et al. Longitudinal assessment of serum 25-hydroxyvitamin D levels during pregnancy and postpartum. Nutrients 2023; 15: 2012.",
    "15. Ioannou C, Javaid MK, Mahon P, et al. The effect of maternal vitamin D concentration on fetal bone. J Clin Endocrinol Metab 2012; 97: E2070\u20137.",
    "16. Page MJ, McKenzie JE, Bossuyt PM, et al. The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ 2021; 372: n71.",
    "17. McGowan J, Sampson M, Salzwedel DM, et al. PRESS peer review of electronic search strategies: 2015 guideline statement. J Clin Epidemiol 2016; 75: 40\u20136.",
    "18. Aromataris E. JBI manual for evidence synthesis. Adelaide: JBI, 2024.",
    "19. Wells GA, Shea B, O\u2019Connell D, et al. The Newcastle-Ottawa Scale (NOS) for assessing the quality of nonrandomised studies in meta-analyses. Ottawa: Ottawa Hospital Research Institute, 2021.",
    "20. Sterne JAC, Savovi\u0107 J, Page MJ, et al. RoB 2: a revised tool for assessing risk of bias in randomised trials. BMJ 2019; 366: l4898.",
    "21. Guyatt GH, Oxman AD, Vist GE, et al. GRADE: an emerging consensus on rating quality of evidence and strength of recommendations. BMJ 2008; 336: 924\u20136.",
    "22. Santamaria C, Bi WG, Leduc L, et al. Prenatal vitamin D status and offspring\u2019s growth, adiposity and metabolic health: a systematic review and meta-analysis. Br J Nutr 2018; 119: 310\u201319.",
    "23. Bi WG, Nuyt AM, Bhatt DL, et al. Effect of vitamin D supplementation during pregnancy on newborn and infant outcomes: a systematic review and meta-analysis. JAMA Pediatr 2018; 172: 635\u201345.",
    "24. Kovacs CS. Bone development and mineral homeostasis in the fetus and neonate: roles of the calciotropic and phosphotropic hormones. Physiol Rev 2014; 94: 1143\u20131218.",
    "25. Liu NQ, Hewison M. Vitamin D, the placenta and pregnancy. Arch Biochem Biophys 2012; 523: 37\u201347.",
    "26. Lai JKC, Lucas RM, Banks E, Ponsonby AL. Variability in vitamin D assays impairs clinical assessment of vitamin D status. Intern Med J 2012; 42: 43\u201350.",
    "27. Vafaei H, Asadi N, Kasraeian M, et al. Positive effect of low dose vitamin D supplementation on growth of fetal bones: a randomized prospective study. Bone 2019; 122: 136\u201342.",
    "28. Ikenoue S, Tamai J, Akita K, et al. Origins of obesity in the womb: fetal adiposity and its determinants. J Obstet Gynaecol Res 2024; 50: 2178\u201382.",
    "29. Vafaei H, Asadi N, Kasraeian M, et al. Positive effect of low dose vitamin D supplementation on growth of fetal bones: a randomized prospective study. Bone 2019; 122: 136\u201342.",
    "30. Srilekha V, Vijayalakshmi B, Reddy IY, Fathima N. Effect of vitamin D supplementation on fetal growth and development in pregnant women. Biomedicine 2021; 41: 821\u20134.",
    "31. Young BE, McNanley TJ, Cooper EM, et al. Maternal vitamin D status and calcium intake interact to affect fetal skeletal growth in utero in pregnant adolescents. Am J Clin Nutr 2012; 95: 1103\u201312.",
    "32. Lee DH, Ryu HM, Han YJ, et al. Effects of serum 25-hydroxy-vitamin D and fetal bone growth during pregnancy. J Bone Metab 2015; 22: 127\u201333.",
    "33. Lee SB, Jung SH, Lee H, et al. Maternal vitamin D deficiency in early pregnancy and perinatal and long-term outcomes. Heliyon 2023; 9: e19367.",
    "34. Akita K, Ikenoue S, Tamai J, et al. Maternal serum 25-hydroxyvitamin D as a possible modulator of fetal adiposity. Int J Mol Sci 2025; 26: 1234.",
    "35. Ge LP, Pan J, Liang M. Correlation analysis of maternal serum folate and 25(OH)D levels with the incidence of fetal growth restriction in patients with preeclampsia. J Matern Fetal Neonatal Med 2024; 37: 2400688.",
    "36. Alimohammadi S, Esna-Ashari F, Beheshti Rooy RS. Relationship of vitamin D serum level with intrauterine growth retardation in pregnant women. Int J Womens Health Reprod Sci 2020; 8: 221\u20136.",
    "37. Judistiani RTD, Madjid TH, Irianti S, et al. Association of first trimester maternal vitamin D, ferritin and hemoglobin level with third trimester fetal biometry. BMC Pregnancy Childbirth 2019; 19: 112.",
    "38. Kwon KW, Lee YH, Yeo MH, et al. Maternal and fetal effects of gestational vitamin D concentration. Healthcare 2023; 11: 2325.",
    "39. Mahfod HS, El Behery MM, Zaitoun MM, El-sayed HS. The relation between vitamin D deficiency and fetal growth restriction in pregnant women. Egypt J Hosp Med 2022; 89: 6167\u201373.",
    "40. Fern\u00e1ndez-Alonso AM, Fiol-Ruiz G, Chedraui P, P\u00e9rez-L\u00f3pez FR. Lack of correlation between first trimester maternal serum 25-hydroxyvitamin D levels and ultrasound measured crown-rump length and nuchal translucency. Arch Gynecol Obstet 2011; 283: 13\u201318.",
    "41. Vestergaard AL, Justesen S, Volqvartz T, et al. Vitamin D insufficiency among Danish pregnant women\u2014prevalence and association with adverse obstetric outcomes. Acta Obstet Gynecol Scand 2021; 100: 2144\u201351.",
]

for ref in refs:
    p = doc.add_paragraph()
    run = p.add_run(ref)
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15

doc.add_page_break()

# =============================================================================
# FIGURE LEGENDS
# =============================================================================
heading("Figure Legends", 1)

legends = [
    ("Figure 1.", " PRISMA 2020 flow diagram of study selection."),
    ("Figure 2.", " Forest plot of the association between maternal vitamin D deficiency and SGA/FGR risk (k = 4; N = 9033). Fixed-effects model; pooled OR = 1\u00b795 (95% CI 1\u00b747\u20132\u00b759); I\u00b2 = 0%. \u2020 Direction inverted from original study to harmonise deficiency as increased risk."),
    ("Figure 3.", " Forest plot of the association between maternal vitamin D deficiency and fetal biometry/EFW outcomes (k = 4; N = 26,542). Fixed-effects model; pooled OR = 1\u00b716 (95% CI 1\u00b709\u20131\u00b723); I\u00b2 = 33%."),
    ("Figure 4.", " Forest plot of RCT continuous outcomes: mean difference in fetal biometry with vitamin D supplementation vs control. Green: Vafaei et al. 2019 (low risk of bias); blue: Srilekha et al. 2021 (high risk of bias)."),
]

for label, text in legends:
    para_mixed([(label, True, False), (text, False, False)])

doc.add_paragraph()

supp_legends = [
    ("Figure S1.", " JBI Critical Appraisal heatmap (Q1\u2013Q9) and Newcastle\u2013Ottawa Scale bar chart for 24 observational studies."),
    ("Figure S2.", " Leave-one-out sensitivity analysis for the SGA/FGR meta-analysis."),
    ("Figure S3.", " Subgroup forest plot by vitamin D deficiency threshold."),
    ("Figure S4.", " Funnel plot with Duval & Tweedie trim-and-fill analysis."),
    ("Figure S5.", " GRADE evidence profile for all outcomes."),
    ("Figure S6.", " Bubble plot of mean maternal 25(OH)D concentration vs sample size."),
    ("Figure S7.", " Narrative synthesis heatmap (direction of association and quality)."),
    ("Figure S8.", " Study characteristics: design distribution, timeline, geographic map, and quality summary."),
    ("Figure S9.", " Laboratory methods for 25(OH)D measurement and maternal vitamin D concentrations by continent."),
    ("Figure S10.", " Outcome category distribution and vitamin D deficiency threshold summary."),
]

heading("Supplementary Figure Legends", 2)
for label, text in supp_legends:
    para_mixed([(label, True, False), (text, False, False)])

doc.add_page_break()

# =============================================================================
# SUPPLEMENTARY TABLES
# =============================================================================
heading("Supplementary Material", 1)

para("Table S1. Data extraction summary for all 26 included studies (see TableS1_DataExtraction.csv).",
     italic=True)
para("Table S2. Meta-analysis summary statistics (see TableS2_MetaAnalysisSummary.csv).",
     italic=True)
para("Table S3. JBI Critical Appraisal and Newcastle\u2013Ottawa Scale quality scores (see TableS4_JBI_NOS_Quality.csv).",
     italic=True)
para("Table S4. Trim-and-fill analysis results (see TableS_TrimFill.csv).",
     italic=True)
para("Table S5. Leave-one-out sensitivity analysis results (see SR_LOO_Sensitivity.csv).",
     italic=True)

doc.add_paragraph()
para("All supplementary data files are available in the online repository:", italic=True)
para("https://github.com/Audency/maternal-vitamin-d-meta-analysis", bold=True)


# ── SAVE ─────────────────────────────────────────────────────────────────────
out = os.path.expanduser(
    "~/Desktop/AUDENCIO/Producao artigos/Artigo RS Isabel Vitamina D /"
    "Manuscript_VitaminD_SR_Lancet.docx")
doc.save(out)
print(f"Saved: {out}")
