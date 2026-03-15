#!/usr/bin/env python3
"""
Generate Word document with structured abstract for the manuscript.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Styles ───────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = RGBColor(0x11, 0x11, 0x11)
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.5

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Association between maternal vitamin D levels in pregnancy "
    "and fetal growth and adiposity: A systematic review and meta-analysis"
)
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

# ── Authors ──────────────────────────────────────────────────────────────────
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_a = authors.add_run("Isabel O. A., Audencio Victor")
run_a.font.size = Pt(12)
run_a.font.name = 'Times New Roman'

# ── Registration ─────────────────────────────────────────────────────────────
reg = doc.add_paragraph()
reg.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_r = reg.add_run("PROSPERO: CRD42024590338")
run_r.font.size = Pt(11)
run_r.font.italic = True
run_r.font.name = 'Times New Roman'
run_r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacing

# ── ABSTRACT HEADING ─────────────────────────────────────────────────────────
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_h = h.add_run("ABSTRACT")
run_h.bold = True
run_h.font.size = Pt(13)
run_h.font.name = 'Times New Roman'

doc.add_paragraph()  # spacing


def add_section(heading, text):
    p = doc.add_paragraph()
    run_head = p.add_run(heading + ": ")
    run_head.bold = True
    run_head.font.size = Pt(12)
    run_head.font.name = 'Times New Roman'
    run_body = p.add_run(text)
    run_body.font.size = Pt(12)
    run_body.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.5


# ── BACKGROUND ───────────────────────────────────────────────────────────────
add_section("Background",
    "Vitamin D [25(OH)D] plays an essential role in bone metabolism, immune function, "
    "and fetal development. During pregnancy, adequate concentrations of 25(OH)D are "
    "critical for fetal growth, skeletal development, and metabolic programming. "
    "However, the association between maternal vitamin D status and fetal growth outcomes "
    "remains inconsistent across studies, partly due to heterogeneity in study design, "
    "vitamin D deficiency thresholds, and analytical methods."
)

# ── AIM ──────────────────────────────────────────────────────────────────────
add_section("Aim",
    "To synthesize the available scientific evidence on the association between maternal "
    "25(OH)D concentrations during pregnancy and fetal growth and adiposity through "
    "a systematic review with meta-analysis."
)

# ── METHODS ──────────────────────────────────────────────────────────────────
add_section("Methods",
    "A systematic review was conducted according to PRISMA 2020 guidelines and registered "
    "at PROSPERO (CRD42024590338). Six electronic databases (PubMed, EMBASE, Scopus, LILACS, "
    "SciELO, and Web of Science) were searched up to October 2024. Observational studies and "
    "randomized clinical trials (RCTs) evaluating the association between maternal serum "
    "25(OH)D and fetal growth outcomes (small-for-gestational-age [SGA], fetal growth "
    "restriction [FGR], fetal biometry, estimated fetal weight [EFW], and adiposity) were "
    "included. Methodological quality was assessed using the Joanna Briggs Institute (JBI) "
    "Critical Appraisal tools and the Newcastle\u2013Ottawa Scale (NOS) for observational studies, "
    "and the Cochrane Risk of Bias 2 (RoB 2) tool for RCTs. Meta-analyses were performed "
    "using an adaptive model approach: fixed-effects (inverse-variance) when I\u00b2 < 50% and "
    "random-effects (REML) when I\u00b2 \u2265 50%. Odds ratios (OR) were pooled for dichotomous "
    "outcomes and mean differences (MD) for continuous outcomes. Publication bias was evaluated "
    "using Egger\u2019s regression, Begg\u2019s rank test, and Duval & Tweedie trim-and-fill analysis. "
    "Certainty of evidence was assessed using the GRADE framework."
)

# ── RESULTS ──────────────────────────────────────────────────────────────────
add_section("Results",
    "A total of 3,135 records were identified, and 26 studies (24 observational, 2 RCTs) "
    "involving 26,542 participants from 16 countries were included. "
    "For the primary outcome (SGA/FGR), the pooled meta-analysis of four studies (N = 9,033) "
    "demonstrated a significant association between maternal vitamin D deficiency and increased "
    "risk of SGA/FGR (OR = 1.95, 95% CI 1.47\u20132.59; I\u00b2 = 0%; fixed-effects model). "
    "For the secondary outcome (fetal biometry/EFW), four studies (N = 26,542) yielded a pooled "
    "OR of 1.16 (95% CI 1.09\u20131.23; I\u00b2 = 33%; fixed-effects model; GRADE: High). "
    "Two RCTs reported continuous outcomes: Vafaei et al. (2019; low risk of bias) found that "
    "vitamin D supplementation significantly increased femur length (MD = +1.98 mm, 95% CI "
    "1.28\u20132.68) and humerus length (MD = +1.39 mm, 95% CI 0.72\u20132.06). "
    "Leave-one-out sensitivity analyses confirmed the robustness of all pooled estimates. "
    "Trim-and-fill analysis showed no evidence of significant publication bias. "
    "Subgroup analyses by vitamin D deficiency cutoff demonstrated directionally consistent "
    "results across all thresholds. No significant association was found between maternal "
    "25(OH)D and fetal adiposity (k = 2; GRADE: Low)."
)

# ── CONCLUSION ───────────────────────────────────────────────────────────────
add_section("Conclusion",
    "Maternal vitamin D deficiency during pregnancy is significantly associated with impaired "
    "fetal growth, particularly increased risk of SGA/FGR (OR \u2248 2.0) and altered fetal "
    "biometry. Limited RCT evidence suggests that supplementation may improve fetal bone growth. "
    "These findings support routine screening and correction of vitamin D deficiency in pregnancy "
    "as a potential strategy to optimize fetal growth outcomes. Large-scale RCTs with standardized "
    "protocols are needed to confirm causal effects and guide clinical recommendations."
)

# ── KEYWORDS ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
kw = doc.add_paragraph()
run_kw_h = kw.add_run("Keywords: ")
run_kw_h.bold = True
run_kw_h.font.size = Pt(12)
run_kw_h.font.name = 'Times New Roman'
run_kw_b = kw.add_run(
    "vitamin D; 25-hydroxyvitamin D; pregnancy; fetal growth; small-for-gestational-age; "
    "fetal growth restriction; adiposity; meta-analysis; systematic review."
)
run_kw_b.font.size = Pt(12)
run_kw_b.font.name = 'Times New Roman'

# ── WORD COUNT ───────────────────────────────────────────────────────────────
doc.add_paragraph()
wc = doc.add_paragraph()
run_wc = wc.add_run("Word count: ~420")
run_wc.font.size = Pt(10)
run_wc.font.italic = True
run_wc.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run_wc.font.name = 'Times New Roman'

# ── SAVE ─────────────────────────────────────────────────────────────────────
out_path = os.path.expanduser(
    "~/Desktop/AUDENCIO/Producao artigos/Artigo RS Isabel Vitamina D /Abstract_VitaminD_SR.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
